"""Main Telegram Bot for Quiz Management"""
import os
import logging
import io
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, InputFile, Poll
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ContextTypes, ConversationHandler, filters, PollAnswerHandler
)

from test_parser import TestParser
from quiz_storage import QuizStorage

# Optional imports for file handling
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

# Load environment variables
load_dotenv()
TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')

# Setup logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Quiz storage
storage = QuizStorage()

# Conversation states
WAITING_FOR_FILE = 1
WAITING_FOR_TEXT = 2
WAITING_FOR_QUIZ_NAME = 3
WAITING_FOR_QUESTION_EDIT = 4
WAITING_FOR_OPTION_EDIT = 5
WAITING_FOR_CORRECT_OPTION = 6
WAITING_FOR_NEW_QUIZ_NAME = 7
WAITING_FOR_NEW_QUIZ_QUESTION = 8
WAITING_FOR_NEW_QUIZ_OPTIONS = 9
WAITING_FOR_NEW_QUIZ_CORRECT = 10
WAITING_FOR_QUIZ_ACTION = 11

# Store temporary data
user_data = {}

# Store active quizzes for real-time quiz
active_quizzes = {}  # {chat_id: {'quiz': quiz_data, 'current_q': index, 'answers': {user_id: answer_index}}}

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start command handler"""
    keyboard = [
        [InlineKeyboardButton("📄 Test faylini jo'natish", callback_data='upload_test')],
        [InlineKeyboardButton("📝 Test matnini jo'natish", callback_data='text_input')],
        [InlineKeyboardButton("➕ Quiz testlar yaratish", callback_data='create_new_quiz')],
        [InlineKeyboardButton("▶️ Quizni boshlash", callback_data='start_quiz_menu')],
        [InlineKeyboardButton("📋 Mening quizlarim", callback_data='my_quizzes')],
        [InlineKeyboardButton("❓ Qo'llanma", callback_data='help')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        "Assalomu alaykum! 👋\n\n"
        "📚 Quiz Bot-ga xush kelibsiz!\n\n"
        "Men sizga test fayllarini yoki matnni Telegram quizlariga aylantirib beraman "
        "va siz ularni tahrirlashingiz mumkin.",
        reply_markup=reply_markup
    )
    
    return ConversationHandler.END

async def help_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Help command"""
    help_text = (
        "📖 **QO'LLANMA**\n\n"
        "**BOT NIMA QILADI?**\n"
        "Test fayllarini va matnlarni Telegram quizlariga aylantirib beradi.\n\n"
        "**QABUL QILINADIGAN FAYLLAR:**\n"
        "📝 .txt  📕 .pdf  📗 .docx  📘 .doc\n\n"
        "**FORMAT TURLAR:**\n"
        "1️⃣ Q/A: Q: savol? A: javob* \n"
        "2️⃣ Raqamli: 1) savol? a) javob (correct)\n"
        "3️⃣ Sodda: 1 savol? a javob*\n\n"
        "**BUYRUQLAR:**\n"
        "/start - Menyu\n"
        "/newquiz - Yangi quiz\n"
        "/quizzes - Mening quizlar\n"
        "/stop - Quizni to'xtatish\n"
        "/help - Bu qo'llanma"
    )
    
    await update.message.reply_text(help_text, parse_mode='Markdown')

async def newquiz_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start creating new quiz"""
    await update.message.reply_text(
        "📝 **Yangi Quiz Yaratish**\n\n"
        "Quiz uchun nom kiriting:"
    )
    return WAITING_FOR_NEW_QUIZ_NAME

async def quizzes_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show user's quizzes"""
    user_id = update.message.from_user.id
    quizzes = storage.get_user_quizzes(user_id)
    
    if not quizzes:
        await update.message.reply_text("Sizda hali quizlar yo'q.")
        return
    
    text = "📋 **Sizning Quizlaringiz:**\n\n"
    
    for quiz in quizzes[:10]:
        text += f"📌 *{quiz['name']}* ({len(quiz['questions'])} savol)\n"
        text += f"   ID: `{quiz['id']}`\n"
        text += f"   Sana: {quiz['created_at'][:10]}\n\n"
    
    if len(quizzes) > 10:
        text += f"\n... va {len(quizzes) - 10} ta ko'proq\n"
    
    await update.message.reply_text(text, parse_mode='Markdown')

async def quiz_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start a specific quiz by ID"""
    if not context.args:
        await update.message.reply_text(
            "❌ Quiz ID ni ko'rsating!\n"
            "Masalan: `/quiz ID`",
            parse_mode='Markdown'
        )
        return
    
    quiz_id = context.args[0]
    quiz = storage.get_quiz(quiz_id)
    
    if not quiz:
        await update.message.reply_text(f"❌ Quiz topilmadi: `{quiz_id}`", parse_mode='Markdown')
        return
    
    chat_id = update.message.chat_id
    active_quizzes[chat_id] = {
        'quiz': quiz,
        'current_q': 0,
        'answers': {}
    }
    
    await update.message.reply_text(f"▶️ Quiz boshlanmoqda: *{quiz['name']}*", parse_mode='Markdown')
    await show_quiz_question(context, chat_id, quiz)

async def stop_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Stop active quiz"""
    chat_id = update.message.chat_id
    
    if chat_id in active_quizzes:
        del active_quizzes[chat_id]
        await update.message.reply_text("⏹️ Quiz to'xtatildi.")
    else:
        await update.message.reply_text("Hozir faol quiz yo'q.")

async def show_quiz_question(context, chat_id, quiz, show_poll=True):
    """Show current quiz question using Telegram polls"""
    if chat_id not in active_quizzes:
        return
    
    quiz_data = active_quizzes[chat_id]
    q_index = quiz_data['current_q']
    
    if q_index >= len(quiz['questions']):
        # Quiz complete
        await show_quiz_results_final(context, chat_id, quiz_data)
        return
    
    question = quiz['questions'][q_index]
    
    # Create poll using Telegram's native poll feature
    poll_question = question['question']
    poll_options = question['options']
    
    try:
        # Send anonymous poll
        poll_msg = await context.bot.send_poll(
            chat_id=chat_id,
            question=poll_question,
            options=poll_options,
            is_anonymous=True,
            type=Poll.REGULAR,
            open_period=15  # Poll stays open for 15 seconds
        )
        
        # Store poll message ID with question info for answer tracking
        if 'poll_messages' not in quiz_data:
            quiz_data['poll_messages'] = {}
        
        quiz_data['poll_messages'][poll_msg.message_id] = {
            'question_index': q_index,
            'correct_option': question['correct_option_id']
        }
        
        # Auto-advance after poll closes (add 2 seconds buffer)
        context.job_queue.run_once(
            advance_quiz_question,
            when=17,
            data={'chat_id': chat_id, 'quiz_id': quiz['id'], 'question_index': q_index},
            name=f"advance_quiz_{chat_id}_{q_index}"
        )
        
    except Exception as e:
        logger.error(f"Poll sending error: {e}")
        # Fallback to inline keyboard if poll fails
        keyboard = []
        for i, option in enumerate(question['options']):
            keyboard.append([InlineKeyboardButton(
                f"{chr(97+i)}) {option}",
                callback_data=f"q_answer_{q_index}_{i}"
            )])
        
        text = f"❓ <b>Savol {q_index + 1}/{len(quiz['questions'])}</b>\n\n"
        text += f"{question['question']}\n\n"
        
        await context.bot.send_message(
            chat_id=chat_id,
            text=text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )

async def advance_quiz_question(context: ContextTypes.DEFAULT_TYPE):
    """Auto-advance to next question after poll closes"""
    job = context.job
    chat_id = job.data['chat_id']
    question_index = job.data.get('question_index', 0)
    
    if chat_id not in active_quizzes:
        return
    
    quiz_data = active_quizzes[chat_id]
    quiz = quiz_data['quiz']
    
    # Show correct answer
    if question_index < len(quiz['questions']):
        question = quiz['questions'][question_index]
        correct_option = question['options'][question['correct_option_id']]
        
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"✓ <b>To'g'ri javob:</b> {correct_option}",
            parse_mode='HTML'
        )
    
    # Move to next question
    quiz_data['current_q'] += 1
    
    if quiz_data['current_q'] >= len(quiz['questions']):
        await show_quiz_results_final(context, chat_id, quiz_data)
    else:
        await show_quiz_question(context, chat_id, quiz)

async def poll_answer_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle poll answers from users"""
    if update.poll_answer:
        poll_answer = update.poll_answer
        user_id = poll_answer.user_id
        selected_options = poll_answer.option_ids
        
        if not selected_options:
            return
        
        selected_option = selected_options[0]  # Get first (and usually only) selected option
        
        # Find the quiz that has this poll
        # Since poll_answer doesn't include chat_id, we need to search through active_quizzes
        for chat_id, quiz_data in list(active_quizzes.items()):
            if 'poll_messages' not in quiz_data:
                continue
                
            # Look for this poll in the current quiz
            for poll_msg_id, poll_info in quiz_data['poll_messages'].items():
                # Since we can't directly match poll_id, we'll record answer for the current question
                q_index = quiz_data['current_q'] - 1  # Previous question (just answered)
                
                if q_index >= 0 and q_index < len(quiz_data['quiz']['questions']):
                    if user_id not in quiz_data['answers']:
                        quiz_data['answers'][user_id] = {}
                    
                    quiz_data['answers'][user_id][q_index] = selected_option
                    logger.info(f"Recorded answer: user={user_id}, chat={chat_id}, question={q_index}, answer={selected_option}")
                    break

async def show_quiz_results_final(context, chat_id, quiz_data):
    """Show final quiz results"""
    quiz = quiz_data['quiz']
    
    text = f"🏆 <b>{quiz['name']} - Yakuniy Natijalar</b>\n\n"
    text += f"📊 <b>Savollar soni:</b> {len(quiz['questions'])}\n"
    text += f"✓ Quiz tugallandi! Rahmat qatnashganingiz uchun!\n"
    
    await context.bot.send_message(
        chat_id=chat_id,
        text=text,
        parse_mode='HTML'
    )
    
    # Clean up
    if chat_id in active_quizzes:
        del active_quizzes[chat_id]

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle button clicks"""
    query = update.callback_query
    await query.answer()
    
    user_id = query.from_user.id
    
    if query.data == 'upload_test':
        await query.edit_message_text(
            "📄 Test faylini jo'nating:\n\n"
            "🟢 Qo'llab-quvvatlangan formatlar:\n"
            "📝 .txt - Oddiy matn fayli\n"
            "📕 .pdf - PDF fayli\n"
            "📗 .docx - Word fayli (yangi)\n"
            "📘 .doc - Word fayli (eski)"
        )
        return WAITING_FOR_FILE
    
    elif query.data == 'text_input':
        await query.edit_message_text(
            "📝 Test matnini quyidagi formatda jo'nating:\n\n"
            "**Q/A Format:**\n"
            "Q: Savol matni?\n"
            "A: Javob 1\n"
            "A: To'g'ri javob*\n"
            "A: Javob 3\n\n"
            "yoki\n\n"
            "**Raqamli Format:**\n"
            "1) Savol?\n"
            "a) Javob 1\n"
            "b) To'g'ri javob (correct)\n"
            "c) Javob 3"
        )
        return WAITING_FOR_TEXT
    
    elif query.data == 'create_new_quiz':
        await query.edit_message_text(
            "📝 **Yangi Quiz Yaratish**\n\n"
            "Quiz uchun nom kiriting:"
        )
        return WAITING_FOR_NEW_QUIZ_NAME
    
    elif query.data == 'help':
        help_text = (
            "📖 **QO'LLANMA**\n\n"
            "**BOT NIMA QILADI?**\n"
            "Test fayllarini va matnlarni Telegram quizlariga aylantirib beradi.\n\n"
            "**QABUL QILINADIGAN FAYLLAR:**\n"
            "📝 .txt  📕 .pdf  📗 .docx  📘 .doc\n\n"
            "**FORMAT TURLAR:**\n"
            "1️⃣ Q/A: Q: savol? A: javob* \n"
            "2️⃣ Raqamli: 1) savol? a) javob (correct)\n"
            "3️⃣ Sodda: 1 savol? a javob*\n\n"
            "**BUYRUQLAR:**\n"
            "/start - Menyu\n"
            "/newquiz - Yangi quiz\n"
            "/quizzes - Mening quizlar\n"
            "/stop - Quizni to'xtatish\n"
            "/help - Bu qo'llanma"
        )
        
        keyboard = [
            [InlineKeyboardButton("◀️ Orqaga", callback_data='back_menu')]
        ]
        
        await query.edit_message_text(
            help_text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return ConversationHandler.END
    
    elif query.data == 'back_menu':
        keyboard = [
            [InlineKeyboardButton("📄 Test faylini jo'natish", callback_data='upload_test')],
            [InlineKeyboardButton("📝 Test matnini jo'natish", callback_data='text_input')],
            [InlineKeyboardButton("➕ Quiz testlar yaratish", callback_data='create_new_quiz')],
            [InlineKeyboardButton("▶️ Quizni boshlash", callback_data='start_quiz_menu')],
            [InlineKeyboardButton("📋 Mening quizlarim", callback_data='my_quizzes')],
            [InlineKeyboardButton("❓ Qo'llanma", callback_data='help')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.edit_message_text(
            "👋 Quiz Bot-ga xush kelibsiz!",
            reply_markup=reply_markup
        )
        return ConversationHandler.END
    
    elif query.data == 'confirm_save_quiz':
        # Save the quiz
        user_id = query.from_user.id
        quiz_name = context.user_data.get('quiz_name')
        quiz_questions = context.user_data.get('quiz_questions', [])
        
        if not quiz_name or not quiz_questions:
            await query.edit_message_text("❌ Xatolik bor'di. Qayta urinib ko'ring.")
            return ConversationHandler.END
        
        # Save quiz
        quiz_id = storage.save_quiz(quiz_name, quiz_questions, user_id)
        
        # Clear temporary data
        if user_id in user_data:
            del user_data[user_id]
        
        # Show success message
        text = f"✅ <b>Quiz saqlandi!</b>\n\n"
        text += f"📌 <b>Nom:</b> {quiz_name}\n"
        text += f"📊 <b>Savollar:</b> {len(quiz_questions)}\n"
        
        keyboard = [
            [InlineKeyboardButton("▶️ Boshlash", callback_data=f'start_quiz_{quiz_id}')],
            [InlineKeyboardButton("✏️ Tahrirlash", callback_data=f'edit_{quiz_id}')],
            [InlineKeyboardButton("📋 Mening quizlarim", callback_data='my_quizzes')],
            [InlineKeyboardButton("📄 Yangi test", callback_data='upload_test')]
        ]
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )
        
        return ConversationHandler.END
    
    elif query.data == 'edit_before_save':
        # Show questions for editing before saving
        quiz_questions = context.user_data.get('quiz_questions', [])
        
        text = "<b>✏️ Tahrirlash uchun savol tanlang:</b>\n\n"
        keyboard = []
        
        for i, q in enumerate(quiz_questions):
            text += f"{i+1}. {q['question']}\n"
            keyboard.append([InlineKeyboardButton(
                f"✏️ Savol {i+1}",
                callback_data=f'edit_q_before_{i}'
            )])
        
        keyboard.append([InlineKeyboardButton("✅ Saqlash", callback_data='confirm_save_quiz')])
        keyboard.append([InlineKeyboardButton("❌ Bekor qilish", callback_data='cancel_quiz')])
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )
        
        return ConversationHandler.END
    
    elif query.data == 'cancel_quiz':
        # Cancel and go back to menu
        user_id = query.from_user.id
        if user_id in user_data:
            del user_data[user_id]
        
        keyboard = [
            [InlineKeyboardButton("📄 Test faylini jo'natish", callback_data='upload_test')],
            [InlineKeyboardButton("📝 Test matnini jo'natish", callback_data='text_input')],
            [InlineKeyboardButton("➕ Quiz testlar yaratish", callback_data='create_new_quiz')],
            [InlineKeyboardButton("▶️ Quizni boshlash", callback_data='start_quiz_menu')],
            [InlineKeyboardButton("📋 Mening quizlarim", callback_data='my_quizzes')],
            [InlineKeyboardButton("❓ Qo'llanma", callback_data='help')]
        ]
        
        await query.edit_message_text(
            "❌ Bekor qilindi.\n\n👋 Quiz Bot-ga xush kelibsiz!",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )
        
        return ConversationHandler.END
    
    elif query.data == 'add_more_question':
        await query.edit_message_text(
            "📌 Keyingi savolni kiriting:"
        )
        return WAITING_FOR_NEW_QUIZ_QUESTION
    
    elif query.data == 'save_new_quiz':
        user_id = query.from_user.id
        quiz_name = context.user_data.get('new_quiz_name')
        questions = context.user_data.get('new_quiz_questions', [])
        
        if not questions:
            await query.edit_message_text("❌ Kamida 1 ta savol bo'lishi kerak!")
            return ConversationHandler.END
        
        # Save quiz
        quiz_id = storage.save_quiz(quiz_name, questions, user_id)
        
        # Clear data
        context.user_data['new_quiz_name'] = None
        context.user_data['new_quiz_questions'] = []
        
        keyboard = [
            [InlineKeyboardButton("✏️ Tahrirlash", callback_data=f'edit_{quiz_id}')],
            [InlineKeyboardButton("▶️ Boshlash", callback_data=f'start_quiz_{quiz_id}')],
            [InlineKeyboardButton("📋 Mening quizlarim", callback_data='my_quizzes')]
        ]
        
        await query.edit_message_text(
            f"✅ **Quiz tayyor!**\n\n"
            f"📌 Nom: {quiz_name}\n"
            f"📊 Savollar: {len(questions)}",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        
        return ConversationHandler.END
    
    elif query.data.startswith('start_quiz_'):
        quiz_id = query.data.replace('start_quiz_', '')
        quiz = storage.get_quiz(quiz_id)
        
        if not quiz:
            await query.edit_message_text("Quiz topilmadi!")
            return ConversationHandler.END
        
        chat_id = query.message.chat_id
        
        # Initialize quiz
        active_quizzes[chat_id] = {
            'quiz': quiz,
            'current_q': 0,
            'answers': {}
        }
        
        await query.edit_message_text("Quiz boshlanmoqda...")
        
        # Show first question
        await show_quiz_question(context, chat_id, quiz)
        return ConversationHandler.END
    
    elif query.data == 'start_quiz_menu':
        quizzes = storage.get_user_quizzes(user_id)
        
        if not quizzes:
            await query.edit_message_text("Sizda hali quizlar yo'q.\n\nAvval quiz yaratishingiz kerak!")
            return ConversationHandler.END
        
        text = "▶️ <b>Quizni boshlash</b>\n\n"
        text += "<b>Qaysi quizni boshlaylik?</b>\n\n"
        keyboard = []
        
        for quiz in quizzes[:10]:  # Show last 10
            text += f"📌 {quiz['name']} ({len(quiz['questions'])} savol)\n"
            keyboard.append([InlineKeyboardButton(
                f"▶️ {quiz['name']}",
                callback_data=f"start_quiz_{quiz['id']}"
            )])
        
        keyboard.append([InlineKeyboardButton("◀️ Orqaga", callback_data='back_menu')])
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )
        return ConversationHandler.END
    
    elif query.data == 'my_quizzes':
        quizzes = storage.get_user_quizzes(user_id)
        
        if not quizzes:
            await query.edit_message_text("Sizda hali quizlar yo'q.")
            return ConversationHandler.END
        
        text = "📋 **Sizning Quizlaringiz:**\n\n"
        keyboard = []
        
        for quiz in quizzes[:10]:  # Show last 10
            text += f"📌 {quiz['name']} ({len(quiz['questions'])} savol)\n"
            keyboard.append([InlineKeyboardButton(
                f"  {quiz['name']}",
                callback_data=f"quiz_{quiz['id']}"
            )])
        
        keyboard.append([InlineKeyboardButton("◀️ Orqaga", callback_data='back_menu')])
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return ConversationHandler.END
    
    elif query.data.startswith('q_answer_'):
        # Handle quiz answer
        parts = query.data.replace('q_answer_', '').split('_')
        q_index = int(parts[0])
        answer_index = int(parts[1])
        
        user_id = query.from_user.id
        chat_id = query.message.chat_id
        
        if chat_id not in active_quizzes:
            await query.edit_message_text("Quiz topilmadi!")
            return ConversationHandler.END
        
        # Record answer
        if user_id not in active_quizzes[chat_id]['answers']:
            active_quizzes[chat_id]['answers'][user_id] = {}
        
        active_quizzes[chat_id]['answers'][user_id][q_index] = answer_index
        
        # Check if this was the last question
        quiz = active_quizzes[chat_id]['quiz']
        if q_index >= len(quiz['questions']) - 1:
            # Last question - show results immediately
            await show_quiz_results_final(context, chat_id, active_quizzes[chat_id])
        else:
            # More questions - show next question
            active_quizzes[chat_id]['current_q'] = q_index + 1
            await show_quiz_question(context, chat_id, quiz)
        
        return ConversationHandler.END
    
    elif query.data.startswith('q_next_'):
        # Go to next question
        user_id = query.from_user.id
        chat_id = query.message.chat_id
        
        if chat_id not in active_quizzes:
            await query.edit_message_text("Quiz topilmadi!")
            return ConversationHandler.END
        
        active_quizzes[chat_id]['current_q'] += 1
        quiz = active_quizzes[chat_id]['quiz']
        
        if active_quizzes[chat_id]['current_q'] >= len(quiz['questions']):
            await show_quiz_results_final(context, chat_id, active_quizzes[chat_id])
        else:
            await show_quiz_question(context, chat_id, quiz)
        return ConversationHandler.END
    
    elif query.data.startswith('quiz_'):
        quiz_id = query.data.replace('quiz_', '')
        quiz = storage.get_quiz(quiz_id)
        
        if not quiz:
            await query.edit_message_text("Quiz topilmadi.")
            return ConversationHandler.END
        
        text = f"📋 **{quiz['name']}**\n"
        text += f"Savollar: {len(quiz['questions'])}\n"
        text += f"Yaratilgan: {quiz['created_at'][:10]}\n\n"
        
        for i, q in enumerate(quiz['questions'][:5]):
            correct = q['options'][q['correct_option_id']]
            text += f"{i+1}. {q['question']}\n"
            text += f"   ✓ {correct}\n"
        
        if len(quiz['questions']) > 5:
            text += f"\n... va {len(quiz['questions']) - 5} ta ko'proq savol"
        
        keyboard = [
            [InlineKeyboardButton("▶️ Boshlash", callback_data=f'start_quiz_{quiz_id}')],
            [InlineKeyboardButton("✏️ Tahrirlash", callback_data=f'edit_{quiz_id}')],
            [InlineKeyboardButton("🗑️ O'chirish", callback_data=f'delete_{quiz_id}')],
            [InlineKeyboardButton("◀️ Orqaga", callback_data='my_quizzes')]
        ]
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return ConversationHandler.END
    
    elif query.data.startswith('edit_'):
        quiz_id = query.data.replace('edit_', '')
        logger.info(f"Edit quiz: quiz_id={quiz_id}")
        quiz = storage.get_quiz(quiz_id)
        
        if not quiz:
            logger.error(f"Quiz not found: quiz_id={quiz_id}")
            await query.edit_message_text(
                f"❌ Quiz topilmadi!\n\n"
                f"ID: {quiz_id}\n\n"
                f"Mening quizlarim'dan quizni tanlang.",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("📋 Mening quizlarim", callback_data='my_quizzes')],
                    [InlineKeyboardButton("◀️ Orqaga", callback_data='back_menu')]
                ])
            )
            return ConversationHandler.END
        
        context.user_data['current_quiz_id'] = quiz_id
        
        text = f"✏️ **{quiz['name']} - Tahrirlash**\n\n"
        keyboard = []
        
        for i, q in enumerate(quiz['questions']):
            text += f"{i+1}. {q['question']}\n"
            keyboard.append([InlineKeyboardButton(
                f"✏️ Savol {i+1}",
                callback_data=f'edit_q_{i}'
            )])
        
        keyboard.append([InlineKeyboardButton("◀️ Orqaga", callback_data=f"quiz_{quiz_id}")])
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return ConversationHandler.END
    
    elif query.data.startswith('edit_q_'):
        q_index = int(query.data.replace('edit_q_', ''))
        quiz_id = context.user_data.get('current_quiz_id')
        quiz = storage.get_quiz(quiz_id)
        
        context.user_data['edit_question_index'] = q_index
        question = quiz['questions'][q_index]
        
        text = f"Savol {q_index + 1}: {question['question']}\n\n"
        text += "Javoblar:\n"
        
        for i, opt in enumerate(question['options']):
            mark = "✓" if i == question['correct_option_id'] else " "
            text += f"{chr(97+i)}) [{mark}] {opt}\n"
        
        keyboard = [
            [InlineKeyboardButton("✏️ Savol matni", callback_data='edit_text')],
            [InlineKeyboardButton("✏️ Javoblar", callback_data='edit_options')],
            [InlineKeyboardButton("✓ To'g'ri javob", callback_data='edit_correct')],
            [InlineKeyboardButton("◀️ Orqaga", callback_data=f"edit_{quiz_id}")]
        ]
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return ConversationHandler.END
    
    elif query.data == 'edit_text':
        await query.edit_message_text(
            "✏️ Savol matni uchun yangi matnni yozing:"
        )
        context.user_data['edit_mode'] = 'question'
        return WAITING_FOR_QUESTION_EDIT
    
    elif query.data == 'edit_options':
        await query.edit_message_text(
            "✏️ Javoblarni quyidagi formatda yozing (har biri alohida qatordan):\n\n"
            "Javob 1\n"
            "Javob 2\n"
            "Javob 3"
        )
        context.user_data['edit_mode'] = 'options'
        return WAITING_FOR_OPTION_EDIT
    
    elif query.data == 'edit_correct':
        quiz_id = context.user_data.get('current_quiz_id')
        quiz = storage.get_quiz(quiz_id)
        q_index = context.user_data.get('edit_question_index')
        question = quiz['questions'][q_index]
        
        text = "To'g'ri javobni tanlang:\n\n"
        keyboard = []
        
        for i, opt in enumerate(question['options']):
            text += f"{chr(97+i)}) {opt}\n"
            keyboard.append([InlineKeyboardButton(
                f"{chr(97+i).upper()}",
                callback_data=f'set_correct_{i}'
            )])
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return ConversationHandler.END
    
    elif query.data.startswith('set_correct_'):
        correct_index = int(query.data.replace('set_correct_', ''))
        quiz_id = context.user_data.get('current_quiz_id')
        q_index = context.user_data.get('edit_question_index')
        
        storage.update_question(
            quiz_id, q_index,
            new_correct_option=correct_index
        )
        
        await query.edit_message_text("✓ To'g'ri javob saqlandi!")
        
        # Return to edit screen
        quiz = storage.get_quiz(quiz_id)
        question = quiz['questions'][q_index]
        
        text = f"Savol {q_index + 1}: {question['question']}\n\n"
        text += "Javoblar:\n"
        
        for i, opt in enumerate(question['options']):
            mark = "✓" if i == question['correct_option_id'] else " "
            text += f"{chr(97+i)}) [{mark}] {opt}\n"
        
        keyboard = [
            [InlineKeyboardButton("✏️ Savol matni", callback_data='edit_text')],
            [InlineKeyboardButton("✏️ Javoblar", callback_data='edit_options')],
            [InlineKeyboardButton("✓ To'g'ri javob", callback_data='edit_correct')],
            [InlineKeyboardButton("◀️ Orqaga", callback_data=f"edit_{quiz_id}")]
        ]
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return ConversationHandler.END
    
    elif query.data.startswith('delete_'):
        quiz_id = query.data.replace('delete_', '')
        
        keyboard = [
            [InlineKeyboardButton("✓ Ha, o'chirish", callback_data=f'confirm_delete_{quiz_id}')],
            [InlineKeyboardButton("❌ Yo'q", callback_data=f'quiz_{quiz_id}')]
        ]
        
        await query.edit_message_text(
            "Rostdan ham o'chirmoqchisiz?",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return ConversationHandler.END
    
    elif query.data.startswith('confirm_delete_'):
        quiz_id = query.data.replace('confirm_delete_', '')
        storage.delete_quiz(quiz_id)
        
        await query.edit_message_text("✓ Quiz o'chirildi!")
        
        # Show quizzes list
        quizzes = storage.get_user_quizzes(query.from_user.id)
        
        if not quizzes:
            await query.edit_message_text("Sizda hali quizlar yo'q.")
            return ConversationHandler.END
        
        text = "📋 **Sizning Quizlaringiz:**\n\n"
        keyboard = []
        
        for quiz in quizzes[:10]:
            text += f"📌 {quiz['name']} ({len(quiz['questions'])} savol)\n"
            keyboard.append([InlineKeyboardButton(
                f"  {quiz['name']}",
                callback_data=f"quiz_{quiz['id']}"
            )])
        
        keyboard.append([InlineKeyboardButton("◀️ Orqaga", callback_data='back_menu')])
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return ConversationHandler.END
    
    return ConversationHandler.END

async def file_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle file uploads - supports txt, pdf, docx, doc"""
    if update.message.document:
        file = update.message.document
        file_name = file.file_name.lower()
        
        # Check file type
        allowed_types = ('.txt', '.pdf', '.docx', '.doc')
        if not any(file_name.endswith(ext) for ext in allowed_types):
            await update.message.reply_text(
                "❌ Iltimos, quyidagi fayllardan birini jo'nating:\n"
                "📄 .txt (Matn fayli)\n"
                "📕 .pdf (PDF fayli)\n"
                "📗 .docx (Word fayli)\n"
                "📘 .doc (Word fayli)"
            )
            return
        
        try:
            # Download file
            file_obj = await context.bot.get_file(file.file_id)
            file_content = await file_obj.download_as_bytearray()
            
            # Extract text based on file type
            if file_name.endswith('.txt'):
                content = file_content.decode('utf-8', errors='ignore')
            
            elif file_name.endswith('.pdf'):
                if PdfReader is None:
                    await update.message.reply_text("❌ PDF tahlil qilish imkoniyati mavjud emas.")
                    return
                
                try:
                    pdf_file = io.BytesIO(file_content)
                    pdf_reader = PdfReader(pdf_file)
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text()
                except Exception as e:
                    await update.message.reply_text(f"❌ PDF o'qib bo'lmadi: {str(e)}")
                    return
            
            elif file_name.endswith(('.docx', '.doc')):
                if Document is None:
                    await update.message.reply_text("❌ Word fayl tahlil qilish imkoniyati mavjud emas.")
                    return
                
                try:
                    docx_file = io.BytesIO(file_content)
                    doc = Document(docx_file)
                    
                    # Extract from paragraphs
                    content = "\n".join([para.text for para in doc.paragraphs])
                    
                    # Also extract from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content += "\n" + cell.text
                    
                    logger.info(f"DOCX extracted {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, content length={len(content)}")
                except Exception as e:
                    await update.message.reply_text(f"❌ Word fayl o'qib bo'lmadi: {str(e)}")
                    return
            
            else:
                content = file_content.decode('utf-8', errors='ignore')
            
            # Parse test
            questions, format_type = TestParser.parse(content)
            
            if not questions:
                await update.message.reply_text(
                    "❌ Fayldan savol topilmadi.\n\n"
                    "Lütfen, faylning formatini tekshiring:\n"
                    "- Q/A format\n"
                    "- Raqamli format"
                )
                return
            
            # Store temporary data
            user_id = update.message.from_user.id
            user_data[user_id] = {
                'questions': questions,
                'format': format_type
            }
            
            # Ask for quiz name
            text = f"✓ {len(questions)} ta savol topildi! ({format_type})\n\n"
            text += "Iltimos, quiz nomini kiriting:"
            
            await update.message.reply_text(text)
            return WAITING_FOR_QUIZ_NAME
        
        except Exception as e:
            await update.message.reply_text(f"❌ Xatolik: {str(e)}")
            logger.error(f"File processing error: {str(e)}")
            return WAITING_FOR_FILE

async def text_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle plain text quiz input"""
    text = update.message.text.strip()
    
    if not text or len(text) < 5:
        await update.message.reply_text(
            "❌ Test matni juda qisqa.\n\n"
            "Iltimos, Q/A yoki raqamli formatda matn jo'nating."
        )
        return WAITING_FOR_TEXT
    
    try:
        # Parse test
        questions, format_type = TestParser.parse(text)
        
        if not questions:
            await update.message.reply_text(
                "❌ Matndan savol topilmadi.\n\n"
                "Formatni tekshiring:\n\n"
                "**Q/A Format:**\n"
                "Q: Savol?\n"
                "A: Javob 1\n"
                "A: To'g'ri javob*\n\n"
                "**Raqamli Format:**\n"
                "1) Savol?\n"
                "a) Javob 1\n"
                "b) To'g'ri javob (correct)"
            )
            return WAITING_FOR_TEXT
        
        # Store temporary data
        user_id = update.message.from_user.id
        user_data[user_id] = {
            'questions': questions,
            'format': format_type
        }
        
        # Ask for quiz name
        text = f"✓ {len(questions)} ta savol topildi! ({format_type})\n\n"
        text += "Iltimos, quiz nomini kiriting:"
        
        await update.message.reply_text(text)
        return WAITING_FOR_QUIZ_NAME
    
    except Exception as e:
        await update.message.reply_text(f"❌ Xatolik: {str(e)}")
        logger.error(f"Text parsing error: {str(e)}")
        return WAITING_FOR_TEXT

async def quiz_name_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle quiz name input"""
    quiz_name = update.message.text.strip()
    
    if not quiz_name or len(quiz_name) < 1:
        await update.message.reply_text("❌ Nomni kiriting!")
        return WAITING_FOR_QUIZ_NAME
    
    user_id = update.message.from_user.id
    quiz_data = user_data.get(user_id)
    
    if not quiz_data:
        await update.message.reply_text("❌ Xatolik bor'di. Qayta urinib ko'ring.")
        return ConversationHandler.END
    
    # Store quiz name temporarily for preview
    context.user_data['quiz_name'] = quiz_name
    context.user_data['quiz_questions'] = quiz_data['questions']
    
    # Show preview of questions
    text = f"<b>📋 Quiz Preview</b>\n\n"
    text += f"<b>Nom:</b> {quiz_name}\n"
    text += f"<b>Savollar:</b> {len(quiz_data['questions'])}\n\n"
    
    text += "<b>Savollar ro'yxati:</b>\n"
    for i, q in enumerate(quiz_data['questions'][:5], 1):  # Show first 5
        correct_opt = q['options'][q.get('correct_option_id', 0)]
        text += f"{i}. {q['question']}\n"
        text += f"   ✓ {correct_opt}\n"
    
    if len(quiz_data['questions']) > 5:
        text += f"\n... va {len(quiz_data['questions']) - 5} ta ko'proq savol"
    
    keyboard = [
        [InlineKeyboardButton("✅ Saqlash", callback_data='confirm_save_quiz')],
        [InlineKeyboardButton("📝 Tahrirlash", callback_data='edit_before_save')],
        [InlineKeyboardButton("❌ Bekor qilish", callback_data='cancel_quiz')]
    ]
    
    await update.message.reply_text(
        text,
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='HTML'
    )
    
    return ConversationHandler.END

async def question_edit_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle question text editing"""
    new_text = update.message.text.strip()
    quiz_id = context.user_data.get('current_quiz_id')
    q_index = context.user_data.get('edit_question_index')
    
    storage.update_question(quiz_id, q_index, new_question=new_text)
    
    await update.message.reply_text("✓ Savol yangilandi!")
    
    # Return to edit screen
    quiz = storage.get_quiz(quiz_id)
    question = quiz['questions'][q_index]
    
    text = f"Savol {q_index + 1}: {question['question']}\n\n"
    text += "Javoblar:\n"
    
    for i, opt in enumerate(question['options']):
        mark = "✓" if i == question['correct_option_id'] else " "
        text += f"{chr(97+i)}) [{mark}] {opt}\n"
    
    keyboard = [
        [InlineKeyboardButton("✏️ Savol matni", callback_data='edit_text')],
        [InlineKeyboardButton("✏️ Javoblar", callback_data='edit_options')],
        [InlineKeyboardButton("✓ To'g'ri javob", callback_data='edit_correct')],
        [InlineKeyboardButton("◀️ Orqaga", callback_data=f"edit_{quiz_id}")]
    ]
    
    await update.message.reply_text(
        text,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    return ConversationHandler.END

async def options_edit_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle options editing"""
    options_text = update.message.text.strip()
    new_options = [opt.strip() for opt in options_text.split('\n') if opt.strip()]
    
    if len(new_options) < 2:
        await update.message.reply_text("❌ Kamida 2 ta javob bo'lishi kerak!")
        return WAITING_FOR_OPTION_EDIT
    
    quiz_id = context.user_data.get('current_quiz_id')
    q_index = context.user_data.get('edit_question_index')
    quiz = storage.get_quiz(quiz_id)
    
    # Keep correct option if still valid
    correct_option = min(quiz['questions'][q_index]['correct_option_id'], len(new_options) - 1)
    
    storage.update_question(
        quiz_id, q_index,
        new_options=new_options,
        new_correct_option=correct_option
    )
    
    await update.message.reply_text("✓ Javoblar yangilandi!")
    
    # Return to edit screen
    quiz = storage.get_quiz(quiz_id)
    question = quiz['questions'][q_index]
    
    text = f"Savol {q_index + 1}: {question['question']}\n\n"
    text += "Javoblar:\n"
    
    for i, opt in enumerate(question['options']):
        mark = "✓" if i == question['correct_option_id'] else " "
        text += f"{chr(97+i)}) [{mark}] {opt}\n"
    
    keyboard = [
        [InlineKeyboardButton("✏️ Savol matni", callback_data='edit_text')],
        [InlineKeyboardButton("✏️ Javoblar", callback_data='edit_options')],
        [InlineKeyboardButton("✓ To'g'ri javob", callback_data='edit_correct')],
        [InlineKeyboardButton("◀️ Orqaga", callback_data=f"edit_{quiz_id}")]
    ]
    
    await update.message.reply_text(
        text,
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    return ConversationHandler.END

async def quiz_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start quiz with /quiz ID command"""
    if not context.args:
        await update.message.reply_text(
            "❌ Ishlatish: `/quiz QUIZ_ID`\n\n"
            "Misol: `/quiz 123456_20260416_180000_Kimya`",
            parse_mode='Markdown'
        )
        return
    
    quiz_id = context.args[0]
    quiz = storage.get_quiz(quiz_id)
    
    if not quiz:
        await update.message.reply_text("❌ Quiz topilmadi!")
        return
    
    chat_id = update.message.chat_id
    user_id = update.message.from_user.id
    
    # Initialize quiz
    active_quizzes[chat_id] = {
        'quiz': quiz,
        'current_q': 0,
        'answers': {}
    }
    
    await update.message.reply_text(
        f"🎯 **{quiz['name']}** boshlandi!\n"
        f"📊 Jami savol: {len(quiz['questions'])}\n\n"
        "Tayyor? 👇",
        parse_mode='Markdown'
    )
    
    # Show first question
    await show_quiz_question(context, chat_id, quiz)

async def new_quiz_name_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle new quiz name input"""
    quiz_name = update.message.text.strip()
    
    if not quiz_name or len(quiz_name) < 1:
        await update.message.reply_text("❌ Nomni kiriting!")
        return WAITING_FOR_NEW_QUIZ_NAME
    
    user_id = update.message.from_user.id
    context.user_data['new_quiz_name'] = quiz_name
    context.user_data['new_quiz_questions'] = []
    
    await update.message.reply_text(
        f"✓ Quiz nomi: *{quiz_name}*\n\n"
        "Birinchi savolni kiriting:",
        parse_mode='Markdown'
    )
    
    return WAITING_FOR_NEW_QUIZ_QUESTION

async def new_quiz_question_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle new quiz question input"""
    question_text = update.message.text.strip()
    
    if not question_text:
        await update.message.reply_text("❌ Savol matnini kiriting!")
        return WAITING_FOR_NEW_QUIZ_QUESTION
    
    context.user_data['current_question'] = question_text
    
    await update.message.reply_text(
        f"📌 Savol: *{question_text}*\n\n"
        "Javoblarni quyidagi formatda kiriting (har biri alohida qatordan):\n"
        "```\n"
        "Javob 1\n"
        "Javob 2\n"
        "Javob 3\n"
        "```",
        parse_mode='Markdown'
    )
    
    return WAITING_FOR_NEW_QUIZ_OPTIONS

async def new_quiz_options_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle new quiz options input"""
    try:
        options_text = update.message.text.strip()
        
        # Parse options and find correct answer marked with *
        raw_options = [opt.strip() for opt in options_text.split('\n') if opt.strip()]
        
        logger.info(f"Processing options: {raw_options}")
        
        if len(raw_options) < 2:
            await update.message.reply_text("❌ Kamida 2 ta javob bo'lishi kerak!")
            return WAITING_FOR_NEW_QUIZ_OPTIONS
        
        # Check for * marker and clean options
        options = []
        correct_idx = -1
        
        for i, opt in enumerate(raw_options):
            cleaned = opt.strip()
            # Check if ends with * (case-insensitive check)
            if '*' in cleaned:
                # Find position of * and remove it
                clean_opt = cleaned.replace('*', '').strip()
                options.append(clean_opt)
                if correct_idx == -1:  # First one with * is the answer
                    correct_idx = i
                logger.info(f"Found correct answer at {i}: {clean_opt}")
            else:
                options.append(cleaned)
        
        context.user_data['current_options'] = options
        logger.info(f"Final options: {options}, correct_idx: {correct_idx}")
        
        # If we found a * marker, use that as correct answer
        if correct_idx != -1:
            context.user_data['correct_option_idx'] = correct_idx
            
            # Show confirmation and proceed to next question
            opts_text = "<b>✓ Javoblar:</b>\n"
            for i, opt in enumerate(options):
                mark = "✓" if i == correct_idx else " "
                opts_text += f"{i}) [{mark}] {opt}\n"
            
            keyboard = [
                [InlineKeyboardButton("➕ Savol qo'shish", callback_data='add_more_question')],
                [InlineKeyboardButton("✓ Quizni saqlash", callback_data='save_new_quiz')]
            ]
            
            # Add question directly
            question = {
                'question': context.user_data['current_question'],
                'options': options,
                'correct_option_id': correct_idx
            }
            
            context.user_data['new_quiz_questions'].append(question)
            logger.info(f"Question added: {question}")
            
            await update.message.reply_text(
                opts_text + f"\n<b>✅ Savol #{len(context.user_data['new_quiz_questions'])} qo'shildi!</b>",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode='HTML'
            )
            
            return WAITING_FOR_QUIZ_ACTION
        else:
            # No * marker, ask for correct answer index
            opts_text = "<b>Javoblar:</b>\n"
            for i, opt in enumerate(options):
                opts_text += f"{i}) {opt}\n"
            
            logger.info(f"No * marker found, asking user to select")
            
            await update.message.reply_text(
                opts_text + "\n<b>To'g'ri javob raqamini kiriting (0, 1, 2...):</b>",
                parse_mode='HTML'
            )
            
            return WAITING_FOR_NEW_QUIZ_CORRECT
    
    except Exception as e:
        logger.error(f"Error in new_quiz_options_handler: {e}")
        await update.message.reply_text(f"❌ Xato: {str(e)}")
        return WAITING_FOR_NEW_QUIZ_OPTIONS

async def new_quiz_correct_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle correct option selection"""
    try:
        correct_idx = int(update.message.text.strip())
        options = context.user_data.get('current_options', [])
        
        if correct_idx < 0 or correct_idx >= len(options):
            await update.message.reply_text(f"❌ Raqamni 0 dan {len(options)-1} gacha tanlang!")
            return WAITING_FOR_NEW_QUIZ_CORRECT
        
        # Add question to quiz
        question = {
            'question': context.user_data['current_question'],
            'options': options,
            'correct_option_id': correct_idx
        }
        
        context.user_data['new_quiz_questions'].append(question)
        
        # Ask for next question
        keyboard = [
            [InlineKeyboardButton("➕ Savol qo'shish", callback_data='add_more_question')],
            [InlineKeyboardButton("✓ Quizni saqlash", callback_data='save_new_quiz')]
        ]
        
        await update.message.reply_text(
            f"✓ Savol #{len(context.user_data['new_quiz_questions'])} qo'shildi!\n\n"
            f"📌 {context.user_data['current_question']}\n"
            f"✓ To'g'ri javob: {options[correct_idx]}",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        
        return WAITING_FOR_QUIZ_ACTION
        
    except ValueError:
        await update.message.reply_text("❌ Raqam kiriting!")
        return WAITING_FOR_NEW_QUIZ_CORRECT

def main():
    """Start the bot"""
    if not TOKEN:
        logger.error("TELEGRAM_BOT_TOKEN not found in .env file!")
        return
    
    # Create application
    app = Application.builder().token(TOKEN).build()
    
    # Add command handlers
    app.add_handler(CommandHandler('start', start))
    app.add_handler(CommandHandler('help', help_handler))
    app.add_handler(CommandHandler('newquiz', newquiz_command))
    app.add_handler(CommandHandler('quizzes', quizzes_command))
    app.add_handler(CommandHandler('quiz', quiz_command))
    app.add_handler(CommandHandler('stop', stop_command))
    
    # Conversation handler for quiz creation and editing
    conv_handler = ConversationHandler(
        entry_points=[
            CallbackQueryHandler(button_handler)
        ],
        states={
            WAITING_FOR_FILE: [
                MessageHandler(filters.Document.ALL, file_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_TEXT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, text_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_QUIZ_NAME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, quiz_name_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_QUESTION_EDIT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, question_edit_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_OPTION_EDIT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, options_edit_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_NEW_QUIZ_NAME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, new_quiz_name_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_NEW_QUIZ_QUESTION: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, new_quiz_question_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_NEW_QUIZ_OPTIONS: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, new_quiz_options_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_NEW_QUIZ_CORRECT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, new_quiz_correct_handler),
                CallbackQueryHandler(button_handler)
            ],
            WAITING_FOR_QUIZ_ACTION: [
                CallbackQueryHandler(button_handler)
            ],
        },
        fallbacks=[CommandHandler('start', start), CallbackQueryHandler(button_handler)]
    )
    
    app.add_handler(conv_handler)
    
    # Add poll answer handler
    app.add_handler(PollAnswerHandler(poll_answer_handler))
    
    # Set bot commands
    async def set_commands(app):
        """Set bot commands"""
        from telegram import BotCommand
        commands = [
            BotCommand("/start", "Bosh menyu"),
            BotCommand("/newquiz", "Yangi quiz yaratish"),
            BotCommand("/quizzes", "Mening quizlarim"),
            BotCommand("/quiz", "Quizni ID orqali boshlash"),
            BotCommand("/stop", "Quizni to'xtatish"),
            BotCommand("/help", "Qo'llanma"),
        ]
        await app.bot.set_my_commands(commands)
    
    # Set post_init callback
    async def post_init(app):
        """Post init callback"""
        await set_commands(app)
    
    app.post_init = post_init
    
    # Start polling
    logger.info("Bot started successfully!")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main()